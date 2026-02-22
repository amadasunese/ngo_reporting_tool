from datetime import datetime
from io import BytesIO
from flask import render_template, request, send_file
from sqlalchemy import func
from extensions import db
from models import Project, StrategicObjective, Indicator, Activity, ActivityAttendance
from routes import bp_reports

def _parse_dates(start: str, end: str):
    start_d = datetime.strptime(start, "%Y-%m-%d").date()
    end_d = datetime.strptime(end, "%Y-%m-%d").date()
    return start_d, end_d

def _get_period_data(project_id: int, start_d, end_d):
    project = Project.query.get_or_404(project_id)

    # SO reach summary
    so_rows = (db.session.query(
        StrategicObjective.id,
        StrategicObjective.so_code,
        StrategicObjective.title,
        func.coalesce(func.sum(ActivityAttendance.male_count), 0).label("male"),
        func.coalesce(func.sum(ActivityAttendance.female_count), 0).label("female"),
    )
    .join(Activity, Activity.strategic_objective_id == StrategicObjective.id)
    .join(ActivityAttendance, ActivityAttendance.activity_id == Activity.id)
    .filter(
        StrategicObjective.project_id == project_id,
        Activity.activity_date >= start_d,
        Activity.activity_date <= end_d
    )
    .group_by(StrategicObjective.id)
    .order_by(StrategicObjective.so_code.asc())
    .all())

    so_summary = [{
        "so_code": r.so_code,
        "title": r.title,
        "male": int(r.male),
        "female": int(r.female),
        "total": int(r.male + r.female),
    } for r in so_rows]

    # Indicator reach summary (computed from linked activities)
    ind_rows = (db.session.query(
        Indicator.id,
        Indicator.indicator_code,
        Indicator.statement,
        func.coalesce(func.sum(ActivityAttendance.male_count), 0).label("male"),
        func.coalesce(func.sum(ActivityAttendance.female_count), 0).label("female"),
    )
    .join(StrategicObjective, Indicator.strategic_objective_id == StrategicObjective.id)
    .join(Activity, Activity.indicator_id == Indicator.id)
    .join(ActivityAttendance, ActivityAttendance.activity_id == Activity.id)
    .filter(
        StrategicObjective.project_id == project_id,
        Activity.activity_date >= start_d,
        Activity.activity_date <= end_d
    )
    .group_by(Indicator.id)
    .order_by(Indicator.indicator_code.asc())
    .all())

    ind_summary = [{
        "code": r.indicator_code,
        "statement": r.statement,
        "male": int(r.male),
        "female": int(r.female),
        "total": int(r.male + r.female),
    } for r in ind_rows]

    # Activities list in period + reach (LEFT JOIN attendance)
    act_rows = (db.session.query(
        Activity.id,
        Activity.activity_date,
        Activity.activity_code,
        Activity.title,
        Activity.status,
        Activity.location,
        StrategicObjective.so_code.label("so_code"),
        Indicator.indicator_code.label("indicator_code"),
        func.coalesce(func.sum(ActivityAttendance.male_count), 0).label("male"),
        func.coalesce(func.sum(ActivityAttendance.female_count), 0).label("female"),
    )
    .join(StrategicObjective, Activity.strategic_objective_id == StrategicObjective.id)
    .outerjoin(Indicator, Activity.indicator_id == Indicator.id)
    .outerjoin(ActivityAttendance, ActivityAttendance.activity_id == Activity.id)
    .filter(
        StrategicObjective.project_id == project_id,
        Activity.activity_date >= start_d,
        Activity.activity_date <= end_d
    )
    .group_by(Activity.id, StrategicObjective.so_code, Indicator.indicator_code)
    .order_by(Activity.activity_date.asc())
    .all())

    activities = []
    for r in act_rows:
        male = int(r.male)
        female = int(r.female)
        activities.append({
            "date": str(r.activity_date),
            "code": r.activity_code,
            "title": r.title,
            "so_code": r.so_code,
            "indicator_code": r.indicator_code,
            "status": r.status,
            "location": r.location,
            "male": male,
            "female": female,
            "total": male + female,
        })

    return {
        "project": project,
        "start": start_d,
        "end": end_d,
        "so_summary": so_summary,
        "ind_summary": ind_summary,
        "activities": activities,
    }

@bp_reports.get("/")
def report_home():
    projects = Project.query.order_by(Project.created_at.desc()).all()
    return render_template("reports/home.html", projects=projects)

@bp_reports.get("/period")
def report_period():
    project_id = request.args.get("project_id", type=int)
    start = request.args.get("start")
    end = request.args.get("end")

    projects = Project.query.order_by(Project.created_at.desc()).all()
    if not project_id or not start or not end:
        return render_template("reports/period.html", projects=projects, data=None)

    start_d, end_d = _parse_dates(start, end)
    data = _get_period_data(project_id, start_d, end_d)
    return render_template(
        "reports/period.html",
        projects=projects,
        data=data,
        project_id=project_id,
        start=start,
        end=end
    )

@bp_reports.get("/period/export/docx")
def export_period_docx():
    project_id = request.args.get("project_id", type=int)
    start = request.args.get("start")
    end = request.args.get("end")
    if not project_id or not start or not end:
        projects = Project.query.order_by(Project.created_at.desc()).all()
        return render_template("reports/period.html", projects=projects, data=None)

    start_d, end_d = _parse_dates(start, end)
    data = _get_period_data(project_id, start_d, end_d)

    from docx import Document

    doc = Document()
    doc.add_heading("Period Report", level=1)
    doc.add_paragraph(f"Project: {data['project'].name}")
    doc.add_paragraph(f"Period: {data['start']} to {data['end']}")

    doc.add_heading("Reach by Strategic Objective", level=2)
    if not data["so_summary"]:
        doc.add_paragraph("No attendance found in this period.")
    else:
        t = doc.add_table(rows=1, cols=5)
        hdr = t.rows[0].cells
        hdr[0].text = "SO"
        hdr[1].text = "Title"
        hdr[2].text = "Male"
        hdr[3].text = "Female"
        hdr[4].text = "Total"
        for r in data["so_summary"]:
            row = t.add_row().cells
            row[0].text = r["so_code"]
            row[1].text = r["title"]
            row[2].text = str(r["male"])
            row[3].text = str(r["female"])
            row[4].text = str(r["total"])

    doc.add_heading("Reach by Indicator", level=2)
    if not data["ind_summary"]:
        doc.add_paragraph("No linked-indicator attendance found in this period.")
    else:
        t = doc.add_table(rows=1, cols=5)
        hdr = t.rows[0].cells
        hdr[0].text = "Indicator"
        hdr[1].text = "Statement"
        hdr[2].text = "Male"
        hdr[3].text = "Female"
        hdr[4].text = "Total"
        for r in data["ind_summary"]:
            row = t.add_row().cells
            row[0].text = r["code"]
            row[1].text = r["statement"]
            row[2].text = str(r["male"])
            row[3].text = str(r["female"])
            row[4].text = str(r["total"])

    doc.add_heading("Activities in Period", level=2)
    if not data["activities"]:
        doc.add_paragraph("No activities found in this period.")
    else:
        t = doc.add_table(rows=1, cols=10)
        hdr = t.rows[0].cells
        hdr[0].text = "Date"
        hdr[1].text = "Activity Code"
        hdr[2].text = "Title"
        hdr[3].text = "SO"
        hdr[4].text = "Indicator"
        hdr[5].text = "Status"
        hdr[6].text = "Location"
        hdr[7].text = "Male"
        hdr[8].text = "Female"
        hdr[9].text = "Total"
        for a in data["activities"]:
            row = t.add_row().cells
            row[0].text = a["date"]
            row[1].text = a["code"] or ""
            row[2].text = a["title"]
            row[3].text = a["so_code"]
            row[4].text = a["indicator_code"] or ""
            row[5].text = a["status"]
            row[6].text = a["location"] or ""
            row[7].text = str(a["male"])
            row[8].text = str(a["female"])
            row[9].text = str(a["total"])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = f"period_report_{data['project'].id}_{data['start']}_to_{data['end']}.docx"
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@bp_reports.get("/period/export/xlsx")
def export_period_xlsx():
    project_id = request.args.get("project_id", type=int)
    start = request.args.get("start")
    end = request.args.get("end")
    if not project_id or not start or not end:
        projects = Project.query.order_by(Project.created_at.desc()).all()
        return render_template("reports/period.html", projects=projects, data=None)

    start_d, end_d = _parse_dates(start, end)
    data = _get_period_data(project_id, start_d, end_d)

    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    ws1 = wb.active
    ws1.title = "SO Summary"
    ws1.append(["SO", "Title", "Male", "Female", "Total"])
    for r in data["so_summary"]:
        ws1.append([r["so_code"], r["title"], r["male"], r["female"], r["total"]])

    ws2 = wb.create_sheet("Indicator Summary")
    ws2.append(["Indicator", "Statement", "Male", "Female", "Total"])
    for r in data["ind_summary"]:
        ws2.append([r["code"], r["statement"], r["male"], r["female"], r["total"]])

    ws3 = wb.create_sheet("Activities")
    ws3.append(["Date", "Activity Code", "Title", "SO", "Indicator", "Status", "Location", "Male", "Female", "Total"])
    for a in data["activities"]:
        ws3.append([
            a["date"], a["code"] or "", a["title"], a["so_code"], a["indicator_code"] or "",
            a["status"], a["location"] or "", a["male"], a["female"], a["total"]
        ])

    for ws in (ws1, ws2, ws3):
        for col in range(1, ws.max_column + 1):
            max_len = 0
            col_letter = get_column_letter(col)
            for cell in ws[col_letter]:
                v = "" if cell.value is None else str(cell.value)
                if len(v) > max_len:
                    max_len = len(v)
            ws.column_dimensions[col_letter].width = min(max_len + 2, 60)

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    filename = f"period_report_{data['project'].id}_{data['start']}_to_{data['end']}.xlsx"
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
