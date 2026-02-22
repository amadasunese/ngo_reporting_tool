import os
import uuid
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from io import BytesIO
from datetime import datetime
from flask import render_template, request, redirect, url_for, flash, session, send_file

from flask import current_app
from routes import bp_testscore
from typing import Union

# File handling
ALLOWED_EXTENSIONS = {"csv", "xlsx"}

def _base_dirs():
    # Store uploads and charts inside app root for simplicity
    root = current_app.root_path
    uploads = os.path.join(root, "uploads", "testscore")
    charts = os.path.join(root, "static", "charts", "testscore")
    os.makedirs(uploads, exist_ok=True)
    os.makedirs(charts, exist_ok=True)
    return uploads, charts

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def read_dataset(filepath: str) -> pd.DataFrame:
    ext = filepath.rsplit(".", 1)[1].lower()
    if ext == "csv":
        return pd.read_csv(filepath)
    return pd.read_excel(filepath)

def analyze_data(df: pd.DataFrame, disaggregate: bool = False):
    # Normalize columns
    cols = {c.strip().lower(): c for c in df.columns}
    if "pre_test" not in cols or "post_test" not in cols:
        raise ValueError("Dataset must contain 'pre_test' and 'post_test' columns.")

    pre = df[cols["pre_test"]]
    post = df[cols["post_test"]]

    overall = {
        "n": int(len(df)),
        "mean_pre": float(np.nanmean(pre)),
        "mean_post": float(np.nanmean(post)),
        "gain": float(np.nanmean(post) - np.nanmean(pre)),
        "pct_gain": float(((np.nanmean(post) - np.nanmean(pre)) / np.nanmean(pre)) * 100) if np.nanmean(pre) != 0 else 0.0,  # Fixed condition
    }  # Ensure the dictionary is properly closed

    gender_df = None

    if disaggregate:
        # accept 'gender' or 'gend' columns (as in your analyzer)
        gcol = cols.get("gender") or cols.get("gend")
        if not gcol:
            raise ValueError("To disaggregate by gender, dataset must include 'gender' (or 'gend') column.")

        tmp = df.copy()
        tmp["_gender"] = tmp[gcol].astype(str).str.strip().str.title()
        gender_df = tmp.groupby("_gender").agg(
            n=(" _gender".strip(), "count"),
            mean_pre=(cols["pre_test"], "mean"),
            mean_post=(cols["post_test"], "mean"),
        ).reset_index()

        gender_df["gain"] = gender_df["mean_post"] - gender_df["mean_pre"]
        # Convert to python-native types for safe templating
        gender_df = gender_df.rename(columns={"_gender": "gender"})

    return overall, gender_df

def _save_chart(fig, charts_dir: str, prefix: str) -> str:
    name = f"{prefix}_{uuid.uuid4().hex}.png"
    out = os.path.join(charts_dir, name)
    fig.tight_layout()
    fig.savefig(out, dpi=160)
    plt.close(fig)
    # Return web path for templates
    return f"/static/charts/testscore/{name}"

def generate_chart(mean_pre: float, mean_post: float) -> str:
    uploads, charts = _base_dirs()
    fig = plt.figure(figsize=(5.5, 3.2))
    ax = fig.add_subplot(111)
    ax.bar(["Pre-test", "Post-test"], [mean_pre, mean_post])
    ax.set_title("Average Scores (Overall)")
    ax.set_ylabel("Score")
    return _save_chart(fig, charts, "overall")

def generate_gender_chart(gender_df: pd.DataFrame) -> str:
    uploads, charts = _base_dirs()
    fig = plt.figure(figsize=(6.2, 3.4))
    ax = fig.add_subplot(111)
    labels = gender_df["gender"].tolist()
    pre = gender_df["mean_pre"].tolist()
    post = gender_df["mean_post"].tolist()

    x = np.arange(len(labels))
    w = 0.35
    ax.bar(x - w/2, pre, width=w, label="Pre-test")
    ax.bar(x + w/2, post, width=w, label="Post-test")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("Score")
    ax.set_title("Average Scores by Gender")
    ax.legend()
    return _save_chart(fig, charts, "gender")

# def generate_narrative(overall: dict, gender_df: pd.DataFrame | None) -> str:
def generate_narrative(overall: dict, gender_df: Union[pd.DataFrame, None]) -> str:
    lines = []
    lines.append(f"A total of {overall['n']} participants completed both the pre-test and post-test.")
    lines.append(f"The average pre-test score was {overall['mean_pre']:.2f}, while the average post-test score was {overall['mean_post']:.2f}.")
    lines.append(f"This indicates an average knowledge gain of {overall['gain']:.2f} points (approximately {overall['pct_gain']:.1f}% improvement from baseline).")

    if gender_df is not None and not gender_df.empty:
        parts = []
        for _, r in gender_df.iterrows():
            parts.append(
                f"{r['gender']}: pre {float(r['mean_pre']):.2f} → post {float(r['mean_post']):.2f} (gain {float(r['gain']):.2f})"
            )
        lines.append("Disaggregated by gender, performance changed as follows: " + "; ".join(parts) + ".")
    return "\n".join(lines)

@bp_testscore.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files.get("dataset")
        disaggregate = request.form.get("disaggregate") == "yes"

        if not file or not file.filename or not allowed_file(file.filename):
            flash("Please upload a valid CSV or Excel file.", "error")
            return redirect(url_for("testscore.index"))

        uploads_dir, charts_dir = _base_dirs()
        safe_name = f"{uuid.uuid4().hex}_{file.filename}"
        filepath = os.path.join(uploads_dir, safe_name)
        file.save(filepath)

        try:
            df = read_dataset(filepath)
            overall, gender_df = analyze_data(df, disaggregate=disaggregate)

            overall_chart_path = generate_chart(overall["mean_pre"], overall["mean_post"])
            gender_chart_path = None
            if gender_df is not None and not gender_df.empty:
                gender_chart_path = generate_gender_chart(gender_df)

            narrative = generate_narrative(overall, gender_df)

            # store for export
            session["testscore_overall"] = overall
            session["testscore_gender"] = gender_df.to_dict(orient="records") if gender_df is not None else None
            session["testscore_narrative"] = narrative
            session["testscore_overall_chart"] = overall_chart_path
            session["testscore_gender_chart"] = gender_chart_path

            return render_template(
                "testscore/report.html",
                overall=overall,
                gender_rows=session["testscore_gender"],
                narrative=narrative,
                overall_chart_path=overall_chart_path,
                gender_chart_path=gender_chart_path,
            )

        except Exception as e:
            flash(str(e), "error")
            return redirect(url_for("testscore.index"))

    return render_template("testscore/index.html")

@bp_testscore.post("/export/word")
def export_word():
    overall = session.get("testscore_overall")
    narrative = session.get("testscore_narrative", "")
    gender_rows = session.get("testscore_gender")

    if not overall:
        flash("Nothing to export yet. Run an analysis first.", "error")
        return redirect(url_for("testscore.index"))

    from docx import Document

    doc = Document()
    doc.add_heading("Test Score Analysis Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    doc.add_heading("Overall Summary", level=2)
    doc.add_paragraph(f"Participants (N): {overall['n']}")
    doc.add_paragraph(f"Mean Pre-test: {overall['mean_pre']:.2f}")
    doc.add_paragraph(f"Mean Post-test: {overall['mean_post']:.2f}")
    doc.add_paragraph(f"Mean Gain: {overall['gain']:.2f}")
    doc.add_paragraph(f"% Improvement: {overall['pct_gain']:.1f}%")

    if gender_rows:
        doc.add_heading("Gender Disaggregation", level=2)
        t = doc.add_table(rows=1, cols=5)
        hdr = t.rows[0].cells
        hdr[0].text = "Gender"
        hdr[1].text = "N"
        hdr[2].text = "Mean Pre"
        hdr[3].text = "Mean Post"
        hdr[4].text = "Gain"
        for r in gender_rows:
            row = t.add_row().cells
            row[0].text = str(r.get("gender", ""))
            row[1].text = str(r.get("n", ""))
            row[2].text = f"{float(r.get('mean_pre', 0)):.2f}"
            row[3].text = f"{float(r.get('mean_post', 0)):.2f}"
            row[4].text = f"{float(r.get('gain', 0)):.2f}"

    doc.add_heading("Narrative Interpretation", level=2)
    for line in (narrative or "").split("\n"):
        doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name="test_score_analysis_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

@bp_testscore.post("/export/pdf")
def export_pdf():
    overall = session.get("testscore_overall")
    narrative = session.get("testscore_narrative", "")
    gender_rows = session.get("testscore_gender")

    if not overall:
        flash("Nothing to export yet. Run an analysis first.", "error")
        return redirect(url_for("testscore.index"))

    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Test Score Analysis Report", styles["Title"]))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Overall Summary", styles["Heading2"]))
    story.append(Paragraph(f"Participants (N): {overall['n']}", styles["Normal"]))
    story.append(Paragraph(f"Mean Pre-test: {overall['mean_pre']:.2f}", styles["Normal"]))
    story.append(Paragraph(f"Mean Post-test: {overall['mean_post']:.2f}", styles["Normal"]))
    story.append(Paragraph(f"Mean Gain: {overall['gain']:.2f}", styles["Normal"]))
    story.append(Paragraph(f"% Improvement: {overall['pct_gain']:.1f}%", styles["Normal"]))
    story.append(Spacer(1, 10))

    if gender_rows:
        story.append(Paragraph("Gender Disaggregation", styles["Heading2"]))
        data = [["Gender", "N", "Mean Pre", "Mean Post", "Gain"]]
        for r in gender_rows:
            data.append([
                str(r.get("gender", "")),
                str(r.get("n", "")),
                f"{float(r.get('mean_pre', 0)):.2f}",
                f"{float(r.get('mean_post', 0)):.2f}",
                f"{float(r.get('gain', 0)):.2f}",
            ])
        tbl = Table(data, hAlign="LEFT")
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("PADDING", (0,0), (-1,-1), 6),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 10))

    story.append(Paragraph("Narrative Interpretation", styles["Heading2"]))
    for line in (narrative or "").split("\n"):
        story.append(Paragraph(line, styles["Normal"]))
        story.append(Spacer(1, 4))

    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4)
    doc.build(story)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name="test_score_analysis_report.pdf",
        mimetype="application/pdf",
    )

@bp_testscore.get("/help")
def help_page():
    return render_template("testscore/help.html")

@bp_testscore.get("/about")
def about():
    return render_template("testscore/about.html")

@bp_testscore.get("/manual")
def manual():
    # Use the manual PDF from the original analyzer package (copied into static/manuals/)
    manual_path = os.path.join(current_app.root_path, "static", "manuals", "Test_Score_User_Manual.pdf")
    if not os.path.exists(manual_path):
        flash("Manual not found.", "error")
        return redirect(url_for("testscore.index"))

    return send_file(manual_path, as_attachment=True, download_name="Test_Score_User_Manual.pdf", mimetype="application/pdf")
