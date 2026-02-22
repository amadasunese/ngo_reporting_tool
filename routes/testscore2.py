import os
import uuid
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from io import BytesIO
from datetime import datetime
from flask import render_template, request, redirect, url_for, flash, session, send_file

from flask import current_app
from routes import bp_testscore

# File handling
# ALLOWED_EXTENSIONS = {"csv", "xlsx"}

def _base_dirs():
    # Store uploads and charts inside app root for simplicity
    root = current_app.root_path
    uploads = os.path.join(root, "uploads", "testscore")
    charts = os.path.join(root, "static", "charts", "testscore")
    os.makedirs(uploads, exist_ok=True)
    os.makedirs(charts, exist_ok=True)
    return uploads, charts



import os
import smtplib
from email.message import EmailMessage
import uuid
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg') 
from flask import Flask, render_template, request, send_file, redirect, url_for, flash
from werkzeug.utils import secure_filename

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from flask import session




UPLOAD_FOLDER = "uploads"
CHART_FOLDER = "static/charts"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CHART_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {"csv", "xlsx"}


# ===============================
# CHECK FILE TYPE
# ===============================
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# ===============================
# READ DATASET
# ===============================
def read_dataset(filepath):
    if filepath.endswith(".csv"):
        df = pd.read_csv(filepath)
    else:
        df = pd.read_excel(filepath)

    # normalize headers
    df.columns = df.columns.astype(str).str.strip().str.lower()

    # required columns
    if "pre_test" not in df.columns or "post_test" not in df.columns:
        raise ValueError("Dataset must contain 'pre_test' and 'post_test' columns.")

    return df  # ✅ keep ALL columns (name, gender, class, etc.)


# ===============================
# ANALYSIS
# ===============================
def analyze_data(df, disaggregate=False):
    # Ensure numeric
    df["pre_test"] = pd.to_numeric(df["pre_test"], errors="coerce")
    df["post_test"] = pd.to_numeric(df["post_test"], errors="coerce")
    df = df.dropna(subset=["pre_test", "post_test"])

    if len(df) == 0:
        raise ValueError("No valid data available after filtering.")

    df["gain"] = df["post_test"] - df["pre_test"]

    mean_pre = df["pre_test"].mean()
    mean_post = df["post_test"].mean()
    mean_gain = df["gain"].mean()
    percent_gain = ((mean_post - mean_pre) / mean_pre) * 100 if mean_pre != 0 else 0
    improvement_rate = (df["gain"] > 0).mean() * 100

    overall = {
        "mean_pre": round(mean_pre, 2),
        "mean_post": round(mean_post, 2),
        "mean_gain": round(mean_gain, 2),
        "percent_gain": round(percent_gain, 2),
        "improvement_rate": round(improvement_rate, 2),
    }

    gender_df = None
    if disaggregate and "gender" in df.columns:
        df["gender"] = df["gender"].astype(str).str.strip().str.capitalize()
        df_g = df[df["gender"].isin(["Male", "Female"])]

        if not df_g.empty:
            gender_df = df_g.groupby("gender").agg(
                pre_test=("pre_test", "mean"),
                post_test=("post_test", "mean"),
                gain=("gain", "mean")
            ).reset_index()

            gender_df["percent_gain"] = (
                (gender_df["post_test"] - gender_df["pre_test"]) / gender_df["pre_test"] * 100
            )

            gender_df = gender_df.round(2)

    return overall, gender_df


# ===============================
# GENERATE CHART
# ===============================
def generate_chart(mean_pre, mean_post):
    filename = f"{uuid.uuid4()}.png"
    filepath = os.path.join(CHART_FOLDER, filename)

    plt.figure()
    plt.bar(["Pre-Test", "Post-Test"], [mean_pre, mean_post])
    plt.title("Pre vs Post Test Mean Scores")
    plt.ylabel("Mean Score")
    plt.savefig(filepath)
    plt.close()

    return filepath

def generate_slopegraph(df):
    # Requires participant_id OR name
    id_col = "participant_id" if "participant_id" in df.columns else ("name" if "name" in df.columns else None)
    if not id_col:
        raise ValueError("Slopegraph requires 'participant_id' or 'name' column.")

    tmp = df[[id_col, "pre_test", "post_test"]].copy()
    tmp["pre_test"] = pd.to_numeric(tmp["pre_test"], errors="coerce")
    tmp["post_test"] = pd.to_numeric(tmp["post_test"], errors="coerce")
    tmp = tmp.dropna(subset=["pre_test", "post_test"]).head(40)

    filename = f"{uuid.uuid4()}.png"
    filepath = os.path.join(CHART_FOLDER, filename)

    plt.figure()
    x0, x1 = 0, 1
    for _, r in tmp.iterrows():
        plt.plot([x0, x1], [r["pre_test"], r["post_test"]])
    plt.xticks([0, 1], ["Pre-Test", "Post-Test"])
    plt.ylabel("Score")
    plt.title("Individual Student Journeys (Slopegraph)")
    plt.tight_layout()
    plt.savefig(filepath)
    plt.close()
    return filepath


def generate_grouped_bar_by_class(df):
    # Requires class OR student_class
    class_col = "class" if "class" in df.columns else ("student_class" if "student_class" in df.columns else None)
    if not class_col:
        raise ValueError("Grouped bar requires 'class' or 'student_class' column.")

    tmp = df[[class_col, "pre_test", "post_test"]].copy()
    tmp["pre_test"] = pd.to_numeric(tmp["pre_test"], errors="coerce")
    tmp["post_test"] = pd.to_numeric(tmp["post_test"], errors="coerce")
    tmp = tmp.dropna(subset=["pre_test", "post_test"])

    g = tmp.groupby(class_col).agg(pre_mean=("pre_test", "mean"), post_mean=("post_test", "mean"))
    g["gain_mean"] = g["post_mean"] - g["pre_mean"]
    g = g.reset_index()

    filename = f"{uuid.uuid4()}.png"
    filepath = os.path.join(CHART_FOLDER, filename)

    x = np.arange(len(g[class_col]))
    width = 0.35

    plt.figure()
    plt.bar(x - width/2, g["pre_mean"], width, label="Pre")
    plt.bar(x + width/2, g["post_mean"], width, label="Post")
    plt.xticks(x, g[class_col], rotation=45, ha="right")
    plt.ylabel("Mean Score")
    plt.title("Pre vs Post Mean Scores by Class (Grouped Bar)")
    plt.legend()
    plt.tight_layout()
    plt.savefig(filepath)
    plt.close()
    return filepath


def generate_dumbbell_plot(df):

    df["pre_test"] = pd.to_numeric(df["pre_test"], errors="coerce")
    df["post_test"] = pd.to_numeric(df["post_test"], errors="coerce")

    df = df.dropna(subset=["pre_test", "post_test"])

    if df.empty:
        raise ValueError("No valid numeric values for dumbbell plot.")

    mean_pre = df["pre_test"].mean()
    mean_post = df["post_test"].mean()

    filename = f"{uuid.uuid4()}.png"
    filepath = os.path.join(CHART_FOLDER, filename)

    plt.figure()

    # ✅ Matching dimensions
    x = [mean_pre, mean_post]
    y = [0, 0]

    plt.plot(x, y)
    plt.scatter(x, y)

    plt.yticks([0], ["Overall"])
    plt.xlabel("Mean Score")
    plt.title("Knowledge Gap (Dumbbell Plot)")

    plt.tight_layout()
    plt.savefig(filepath)
    plt.close()

    return filepath



def generate_stacked_gain_metrics(df):
    """
    Looks for column pairs like pre_q1/post_q1, pre_topic_a/post_topic_a, etc.
    Produces a stacked bar chart of GAINS only across metrics.
    """
    cols = [c.lower() for c in df.columns]
    pre_cols = [c for c in cols if c.startswith("pre_") and c not in ("pre_test",)]
    # Pair with post_...
    pairs = []
    for pre in pre_cols:
        suffix = pre.replace("pre_", "", 1)
        post = f"post_{suffix}"
        if post in cols:
            pairs.append((pre, post, suffix))

    if len(pairs) < 2:
        raise ValueError("Stacked gain chart needs multiple metric pairs like pre_q1/post_q1, pre_q2/post_q2 ...")

    gains = []
    labels = []
    for pre, post, label in pairs:
        pre_s = pd.to_numeric(df[pre], errors="coerce")
        post_s = pd.to_numeric(df[post], errors="coerce")
        g = (post_s - pre_s).mean()
        if pd.notna(g):
            gains.append(float(g))
            labels.append(label)

    if len(gains) < 2:
        raise ValueError("Not enough valid numeric metric pairs for stacked gain chart.")

    filename = f"{uuid.uuid4()}.png"
    filepath = os.path.join(CHART_FOLDER, filename)

    plt.figure()
    bottom = 0
    for gain, label in zip(gains, labels):
        plt.bar(["Total Gain"], [gain], bottom=bottom, label=label)
        bottom += gain

    plt.ylabel("Mean Gain")
    plt.title("Stacked Bar (Gain Only) Across Metrics")
    plt.legend()
    plt.tight_layout()
    plt.savefig(filepath)
    plt.close()
    return filepath


# ===============================
# GENERATE NARRATIVE
# ===============================
def generate_narrative(overall, gender_df=None):
    narrative = f"""
Overall analysis shows that the average pre-test score was {overall['mean_pre']},
which increased to {overall['mean_post']} in the post-test.

This reflects an average knowledge gain of {overall['mean_gain']} points,
representing a {overall['percent_gain']}% improvement.

Additionally, {overall['improvement_rate']}% of participants demonstrated measurable improvement.
"""

    if gender_df is not None and not gender_df.empty:

        narrative += "\n\nGender disaggregation indicates the following trends:\n\n"

        for _, row in gender_df.iterrows():
            narrative += (
                f"{row['gender']} participants improved from "
                f"{row['pre_test']} to {row['post_test']}, "
                f"reflecting an average gain of {row['gain']} points "
                f"({row['percent_gain']}% increase).\n\n"
            )

        best = gender_df.loc[gender_df["gain"].idxmax()]
        narrative += (
            f"The highest average improvement was observed among "
            f"{best['gender']} participants."
        )

    return narrative


def generate_gender_chart(gender_df):
    """
    Generates a bar chart comparing Male vs Female mean gain.
    Returns the saved image path.
    """

    filename = f"{uuid.uuid4()}.png"
    filepath = os.path.join(CHART_FOLDER, filename)

    plt.figure()

    plt.bar(gender_df["gender"], gender_df["gain"])

    plt.xlabel("Gender")
    plt.ylabel("Average Knowledge Gain")
    plt.title("Average Knowledge Gain by Gender")

    plt.tight_layout()
    plt.savefig(filepath)
    plt.close()

    return filepath

# ===============================
# EXPORT WORD
# ===============================
def export_word(narrative, overall=None, overall_chart_path=None,
                gender_rows=None, gender_chart_path=None,
                extra_chart_path=None, extra_chart_title=None):

    filename = f"report_{uuid.uuid4()}.docx"
    filepath = os.path.join("static", filename)

    doc = Document()
    doc.add_heading("Pre and Post Test Analysis Report", level=1)

    # Metrics
    if overall:
        doc.add_heading("Key Metrics", level=2)
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        rows = [
            ("Mean Pre-Test", overall.get("mean_pre")),
            ("Mean Post-Test", overall.get("mean_post")),
            ("Mean Gain", overall.get("mean_gain")),
            ("Percent Gain (%)", overall.get("percent_gain")),
            ("Improvement Rate (%)", overall.get("improvement_rate")),
        ]
        # header
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        for i, (k, v) in enumerate(rows, start=1):
            table.cell(i, 0).text = str(k)
            table.cell(i, 1).text = "" if v is None else str(v)

    # Narrative
    doc.add_heading("Narrative Summary", level=2)
    doc.add_paragraph(narrative or "")

    # Charts
    def add_chart(title, chart_path):
        if chart_path:
            abs_path = chart_path
            if not os.path.isabs(abs_path):
                abs_path = os.path.join(os.getcwd(), chart_path)
            if os.path.exists(abs_path):
                doc.add_heading(title, level=2)
                doc.add_picture(abs_path, width=inch * 6)

    add_chart("Overall Chart: Pre vs Post Mean Scores", overall_chart_path)
    add_chart("Gender Chart: Average Knowledge Gain", gender_chart_path)
    if extra_chart_path:
        add_chart(extra_chart_title or "Additional Chart", extra_chart_path)

    # Gender table
    if gender_rows:
        doc.add_heading("Gender Disaggregation", level=2)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        hdr[0].text = "Gender"
        hdr[1].text = "Mean Pre"
        hdr[2].text = "Mean Post"
        hdr[3].text = "Mean Gain"
        hdr[4].text = "% Gain"

        for r in gender_rows:
            row = t.add_row().cells
            row[0].text = str(r[0])
            row[1].text = str(r[1])
            row[2].text = str(r[2])
            row[3].text = str(r[3])
            row[4].text = str(r[4])

    doc.save(filepath)
    return filepath



# ===============================
# EXPORT PDF
# ===============================
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib import colors

def export_pdf(narrative, overall=None, overall_chart_path=None,
               gender_rows=None, gender_chart_path=None,
               extra_chart_path=None, extra_chart_title=None):

    filename = f"report_{uuid.uuid4()}.pdf"
    filepath = os.path.join("static", filename)

    doc = SimpleDocTemplate(filepath)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Pre and Post Test Analysis Report", styles["Heading1"]))
    elements.append(Spacer(1, 0.3 * inch))

    # Metrics
    if overall:
        elements.append(Paragraph("Key Metrics", styles["Heading2"]))
        data = [
            ["Metric", "Value"],
            ["Mean Pre-Test", str(overall.get("mean_pre", ""))],
            ["Mean Post-Test", str(overall.get("mean_post", ""))],
            ["Mean Gain", str(overall.get("mean_gain", ""))],
            ["Percent Gain (%)", str(overall.get("percent_gain", ""))],
            ["Improvement Rate (%)", str(overall.get("improvement_rate", ""))],
        ]
        tbl = Table(data, hAlign="LEFT")
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        elements.append(tbl)
        elements.append(Spacer(1, 0.3 * inch))

    # Narrative
    elements.append(Paragraph("Narrative Summary", styles["Heading2"]))
    elements.append(Paragraph((narrative or "").replace("\n", "<br/>"), styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    # Charts helper
    def add_chart(title, chart_path):
        if chart_path:
            abs_path = chart_path
            if not os.path.isabs(abs_path):
                abs_path = os.path.join(os.getcwd(), chart_path)
            if os.path.exists(abs_path):
                elements.append(Paragraph(title, styles["Heading2"]))
                elements.append(Image(abs_path, width=6 * inch, height=3.5 * inch))
                elements.append(Spacer(1, 0.3 * inch))

    add_chart("Overall Chart: Pre vs Post Mean Scores", overall_chart_path)
    add_chart("Gender Chart: Average Knowledge Gain", gender_chart_path)
    if extra_chart_path:
        add_chart(extra_chart_title or "Additional Chart", extra_chart_path)

    # Gender table
    if gender_rows:
        elements.append(Paragraph("Gender Disaggregation", styles["Heading2"]))
        data = [["Gender", "Mean Pre", "Mean Post", "Mean Gain", "% Gain"]] + [
            [str(r[0]), str(r[1]), str(r[2]), str(r[3]), str(r[4])] for r in gender_rows
        ]
        tbl = Table(data, hAlign="LEFT")
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        elements.append(tbl)

    doc.build(elements)
    return filepath



# def allowed_file(filename: str) -> bool:
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def read_dataset(filepath: str) -> pd.DataFrame:
#     ext = filepath.rsplit(".", 1)[1].lower()
#     if ext == "csv":
#         return pd.read_csv(filepath)
#     return pd.read_excel(filepath)

# def analyze_data(df: pd.DataFrame, disaggregate: bool = False):
#     # Normalize columns
#     cols = {c.strip().lower(): c for c in df.columns}
#     if "pre_test" not in cols or "post_test" not in cols:
#         raise ValueError("Dataset must contain 'pre_test' and 'post_test' columns.")

#     pre = df[cols["pre_test"]]
#     post = df[cols["post_test"]]

#     overall = {
#         "n": int(len(df)),
#         "mean_pre": float(np.nanmean(pre)),
#         "mean_post": float(np.nanmean(post)),
#         "gain": float(np.nanmean(post) - np.nanmean(pre)),
#         "pct_gain": float(((np.nanmean(post) - np.nanmean(pre)) / np.nanmean(pre)) * 100) if np.nanmean(pre) != 0 else 0.0,  # Fixed condition
#     }  # Ensure the dictionary is properly closed

#     gender_df = None

#     if disaggregate:
#         # accept 'gender' or 'gend' columns (as in your analyzer)
#         gcol = cols.get("gender") or cols.get("gend")
#         if not gcol:
#             raise ValueError("To disaggregate by gender, dataset must include 'gender' (or 'gend') column.")

#         tmp = df.copy()
#         tmp["_gender"] = tmp[gcol].astype(str).str.strip().str.title()
#         gender_df = tmp.groupby("_gender").agg(
#             n=(" _gender".strip(), "count"),
#             mean_pre=(cols["pre_test"], "mean"),
#             mean_post=(cols["post_test"], "mean"),
#         ).reset_index()

#         gender_df["gain"] = gender_df["mean_post"] - gender_df["mean_pre"]
#         # Convert to python-native types for safe templating
#         gender_df = gender_df.rename(columns={"_gender": "gender"})

#     return overall, gender_df

# def _save_chart(fig, charts_dir: str, prefix: str) -> str:
#     name = f"{prefix}_{uuid.uuid4().hex}.png"
#     out = os.path.join(charts_dir, name)
#     fig.tight_layout()
#     fig.savefig(out, dpi=160)
#     plt.close(fig)
#     # Return web path for templates
#     return f"/static/charts/testscore/{name}"

# def generate_chart(mean_pre: float, mean_post: float) -> str:
#     uploads, charts = _base_dirs()
#     fig = plt.figure(figsize=(5.5, 3.2))
#     ax = fig.add_subplot(111)
#     ax.bar(["Pre-test", "Post-test"], [mean_pre, mean_post])
#     ax.set_title("Average Scores (Overall)")
#     ax.set_ylabel("Score")
#     return _save_chart(fig, charts, "overall")

# def generate_gender_chart(gender_df: pd.DataFrame) -> str:
#     uploads, charts = _base_dirs()
#     fig = plt.figure(figsize=(6.2, 3.4))
#     ax = fig.add_subplot(111)
#     labels = gender_df["gender"].tolist()
#     pre = gender_df["mean_pre"].tolist()
#     post = gender_df["mean_post"].tolist()

#     x = np.arange(len(labels))
#     w = 0.35
#     ax.bar(x - w/2, pre, width=w, label="Pre-test")
#     ax.bar(x + w/2, post, width=w, label="Post-test")
#     ax.set_xticks(x)
#     ax.set_xticklabels(labels)
#     ax.set_ylabel("Score")
#     ax.set_title("Average Scores by Gender")
#     ax.legend()
#     return _save_chart(fig, charts, "gender")

# def generate_narrative(overall: dict, gender_df: pd.DataFrame | None) -> str:
#     lines = []
#     lines.append(f"A total of {overall['n']} participants completed both the pre-test and post-test.")
#     lines.append(f"The average pre-test score was {overall['mean_pre']:.2f}, while the average post-test score was {overall['mean_post']:.2f}.")
#     lines.append(f"This indicates an average knowledge gain of {overall['gain']:.2f} points (approximately {overall['pct_gain']:.1f}% improvement from baseline).")

#     if gender_df is not None and not gender_df.empty:
#         parts = []
#         for _, r in gender_df.iterrows():
#             parts.append(
#                 f"{r['gender']}: pre {float(r['mean_pre']):.2f} → post {float(r['mean_post']):.2f} (gain {float(r['gain']):.2f})"
#             )
#         lines.append("Disaggregated by gender, performance changed as follows: " + "; ".join(parts) + ".")
#     return "\n".join(lines)



@bp_testscore.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files.get("dataset")
        disaggregate = request.form.get("disaggregate") == "yes"

        if not file or not file.filename or not allowed_file(file.filename):
            flash("Please upload a valid CSV or Excel file.", "error")
            return redirect(url_for("testscore.index"))

        uploads_dir, charts_dir = _base_dirs()
        safe_name = f"{uuid.uuid4().hex}_{file.filename}"
        filepath = os.path.join(uploads_dir, safe_name)
        file.save(filepath)

        try:
            df = read_dataset(filepath)
            overall, gender_df = analyze_data(df, disaggregate=disaggregate)

            overall_chart_path = generate_chart(overall["mean_pre"], overall["mean_post"])
            gender_chart_path = None
            if gender_df is not None and not gender_df.empty:
                gender_chart_path = generate_gender_chart(gender_df)

            narrative = generate_narrative(overall, gender_df)

            # store for export
            session["testscore_overall"] = overall
            session["testscore_gender"] = gender_df.to_dict(orient="records") if gender_df is not None else None
            session["testscore_narrative"] = narrative
            session["testscore_overall_chart"] = overall_chart_path
            session["testscore_gender_chart"] = gender_chart_path

            return render_template(
                "testscore/report.html",
                overall=overall,
                gender_rows=session["testscore_gender"],
                narrative=narrative,
                overall_chart_path=overall_chart_path,
                gender_chart_path=gender_chart_path,
            )

        except Exception as e:
            flash(str(e), "error")
            return redirect(url_for("testscore.index"))

    return render_template("testscore/index.html")

@bp_testscore.post("/export/word")
def export_word():
    overall = session.get("testscore_overall")
    narrative = session.get("testscore_narrative", "")
    gender_rows = session.get("testscore_gender")

    if not overall:
        flash("Nothing to export yet. Run an analysis first.", "error")
        return redirect(url_for("testscore.index"))

    from docx import Document

    doc = Document()
    doc.add_heading("Test Score Analysis Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    doc.add_heading("Overall Summary", level=2)
    doc.add_paragraph(f"Participants (N): {overall['n']}")
    doc.add_paragraph(f"Mean Pre-test: {overall['mean_pre']:.2f}")
    doc.add_paragraph(f"Mean Post-test: {overall['mean_post']:.2f}")
    doc.add_paragraph(f"Mean Gain: {overall['gain']:.2f}")
    doc.add_paragraph(f"% Improvement: {overall['pct_gain']:.1f}%")

    if gender_rows:
        doc.add_heading("Gender Disaggregation", level=2)
        t = doc.add_table(rows=1, cols=5)
        hdr = t.rows[0].cells
        hdr[0].text = "Gender"
        hdr[1].text = "N"
        hdr[2].text = "Mean Pre"
        hdr[3].text = "Mean Post"
        hdr[4].text = "Gain"
        for r in gender_rows:
            row = t.add_row().cells
            row[0].text = str(r.get("gender", ""))
            row[1].text = str(r.get("n", ""))
            row[2].text = f"{float(r.get('mean_pre', 0)):.2f}"
            row[3].text = f"{float(r.get('mean_post', 0)):.2f}"
            row[4].text = f"{float(r.get('gain', 0)):.2f}"

    doc.add_heading("Narrative Interpretation", level=2)
    for line in (narrative or "").split("\n"):
        doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name="test_score_analysis_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

@bp_testscore.post("/export/pdf")
def export_pdf():
    overall = session.get("testscore_overall")
    narrative = session.get("testscore_narrative", "")
    gender_rows = session.get("testscore_gender")

    if not overall:
        flash("Nothing to export yet. Run an analysis first.", "error")
        return redirect(url_for("testscore.index"))

    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Test Score Analysis Report", styles["Title"]))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Overall Summary", styles["Heading2"]))
    story.append(Paragraph(f"Participants (N): {overall['n']}", styles["Normal"]))
    story.append(Paragraph(f"Mean Pre-test: {overall['mean_pre']:.2f}", styles["Normal"]))
    story.append(Paragraph(f"Mean Post-test: {overall['mean_post']:.2f}", styles["Normal"]))
    story.append(Paragraph(f"Mean Gain: {overall['gain']:.2f}", styles["Normal"]))
    story.append(Paragraph(f"% Improvement: {overall['pct_gain']:.1f}%", styles["Normal"]))
    story.append(Spacer(1, 10))

    if gender_rows:
        story.append(Paragraph("Gender Disaggregation", styles["Heading2"]))
        data = [["Gender", "N", "Mean Pre", "Mean Post", "Gain"]]
        for r in gender_rows:
            data.append([
                str(r.get("gender", "")),
                str(r.get("n", "")),
                f"{float(r.get('mean_pre', 0)):.2f}",
                f"{float(r.get('mean_post', 0)):.2f}",
                f"{float(r.get('gain', 0)):.2f}",
            ])
        tbl = Table(data, hAlign="LEFT")
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("PADDING", (0,0), (-1,-1), 6),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 10))

    story.append(Paragraph("Narrative Interpretation", styles["Heading2"]))
    for line in (narrative or "").split("\n"):
        story.append(Paragraph(line, styles["Normal"]))
        story.append(Spacer(1, 4))

    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4)
    doc.build(story)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name="test_score_analysis_report.pdf",
        mimetype="application/pdf",
    )

@bp_testscore.get("/help")
def help_page():
    return render_template("testscore/help.html")

@bp_testscore.get("/about")
def about():
    return render_template("testscore/about.html")

@bp_testscore.get("/manual")
def manual():
    # Use the manual PDF from the original analyzer package (copied into static/manuals/)
    manual_path = os.path.join(current_app.root_path, "static", "manuals", "Test_Score_User_Manual.pdf")
    if not os.path.exists(manual_path):
        flash("Manual not found.", "error")
        return redirect(url_for("testscore.index"))

    return send_file(manual_path, as_attachment=True, download_name="Test_Score_User_Manual.pdf", mimetype="application/pdf")
